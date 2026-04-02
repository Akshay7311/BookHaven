from docx import Document
from docx.shared import Inches
import os

# Paths
artifact_dir = r"C:\Users\aksha\.gemini\antigravity\brain\4711f4dc-fe7c-4ee3-bfeb-44ae8092d02e"
output_file = os.path.join(artifact_dir, "BookHaven_Documentation.docx")

# Create a new Document
doc = Document()

# Title
doc.add_heading('BookHaven User Documentation', 0)

doc.add_paragraph(
    "Welcome to the BookHaven User Guide. This document provides a comprehensive overview "
    "of how to use the BookHaven application, both from a consumer's perspective and an administrator's perspective."
)

# Part 1: Consumer Guide
doc.add_heading('Part 1: Consumer Guide', level=1)

# Home Page
doc.add_heading('1. Home Page', level=2)
doc.add_paragraph('The home page is your gateway to a vast collection of books. You can see featured events, new arrivals, and popular categories.')
doc.add_picture(os.path.join(artifact_dir, 'consumer_home_1774673615018.png'), width=Inches(6))

# Browsing and Searching
doc.add_heading('2. Browsing and Searching', level=2)
doc.add_paragraph("Navigate to the Shop page to browse the entire collection. Use the search bar or category filters to find exactly what you're looking for.")
doc.add_picture(os.path.join(artifact_dir, 'consumer_shop_1774673628957.png'), width=Inches(6))

# Book Details
doc.add_heading('3. Book Details', level=2)
doc.add_paragraph('Click on any book to view detailed information, including its description, price, and gallery images.')
doc.add_picture(os.path.join(artifact_dir, 'consumer_details_1774673644032.png'), width=Inches(6))

# Shopping Cart
doc.add_heading('4. Shopping Cart', level=2)
doc.add_paragraph('Add items to your cart and review them before proceeding to checkout. You can adjust quantities or remove items here.')
doc.add_picture(os.path.join(artifact_dir, 'consumer_cart_1774673662610.png'), width=Inches(6))

# Checkout Process
doc.add_heading('5. Checkout Process', level=2)
doc.add_paragraph('Securely complete your purchase by providing your shipping details and selecting a payment method.')
doc.add_picture(os.path.join(artifact_dir, 'consumer_checkout_1774673793039.png'), width=Inches(6))

# Order Confirmation
doc.add_heading('6. Order Confirmation', level=2)
doc.add_paragraph('Once your order is placed, you will receive a confirmation message and a unique order ID.')
doc.add_picture(os.path.join(artifact_dir, 'consumer_success_1774673844742.png'), width=Inches(6))

# Part 2: Admin Guide
doc.add_heading('Part 2: Admin Guide', level=1)

# Admin Dashboard
doc.add_heading('1. Admin Dashboard', level=2)
doc.add_paragraph('The dashboard provides a high-level overview of store performance, including total revenue, orders, and low-stock alerts.')
doc.add_picture(os.path.join(artifact_dir, 'admin_dashboard_1774673922220.png'), width=Inches(6))

# Managing Books
doc.add_heading('2. Managing Books', level=2)
doc.add_paragraph('Easily add new books or update existing ones. Manage book covers, gallery images, and inventory levels.')
doc.add_picture(os.path.join(artifact_dir, 'admin_books_1774673936254.png'), width=Inches(6))

# Category Management
doc.add_heading('3. Category Management', level=2)
doc.add_paragraph('Organize your catalog by creating and managing categories.')
doc.add_picture(os.path.join(artifact_dir, 'admin_categories_1774673949115.png'), width=Inches(6))

# Order Management
doc.add_heading('4. Order Management', level=2)
doc.add_paragraph('Track and process customer orders. Update order statuses to keep customers informed.')
doc.add_picture(os.path.join(artifact_dir, 'admin_orders_1774673961576.png'), width=Inches(6))

# Save the document
doc.save(output_file)

print(f"Documentation saved to {output_file}")
