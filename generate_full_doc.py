from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# Paths
artifact_dir = r"C:\Users\aksha\.gemini\antigravity\brain\4711f4dc-fe7c-4ee3-bfeb-44ae8092d02e"
output_file = os.path.join(artifact_dir, "BookHaven_Full_Documentation.docx")

# Create a new Document
doc = Document()

# Title Page
title = doc.add_heading('BookHaven - Complete Project Documentation', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph("\n\n")

# 1. Introduction
doc.add_heading('1. Introduction', level=1)
doc.add_paragraph(
    "BookHaven is a premium, enterprise-grade online bookstore built using the MERN stack "
    "(MySQL, Express.js, React, Node.js). It features a robust administration panel, "
    "a feature-rich consumer storefront, and specialized management for categories, banners, and orders."
)

# 2. Technical Architecture
doc.add_heading('2. Technical Architecture', level=1)
doc.add_heading('2.1 Technology Stack', level=2)
doc.add_paragraph("• Frontend: React.js (Vite), Tailwind CSS, React Router.")
doc.add_paragraph("• Backend: Node.js, Express.js.")
doc.add_paragraph("• Database: MySQL with Sequelize ORM.")
doc.add_paragraph("• Security: JWT Authentication, bcrypt hashing, Helmet (HTTP headers).")

# 3. Setup and Installation
doc.add_heading('3. Setup and Installation', level=1)
doc.add_heading('3.1 Backend Configuration', level=2)
doc.add_paragraph("1. Navigate to /server and run: npm install")
doc.add_paragraph("2. Create .env with DB_HOST, DB_USER, DB_PASSWORD, DB_NAME, JWT_SECRET.")
doc.add_paragraph("3. Start: npm run dev")

doc.add_heading('3.2 Frontend Configuration', level=2)
doc.add_paragraph("1. Navigate to /client and run: npm install")
doc.add_paragraph("2. Create .env with VITE_API_URL.")
doc.add_paragraph("3. Start: npm run dev")

# 4. API Reference
doc.add_heading('4. API Reference', level=1)
table = doc.add_table(rows=1, cols=4)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Feature'
hdr_cells[1].text = 'Endpoint'
hdr_cells[2].text = 'Method'
hdr_cells[3].text = 'Description'

api_data = [
    ('Auth', '/auth/login', 'POST', 'Login and get JWT'),
    ('Books', '/books', 'GET', 'Fetch all books'),
    ('Books', '/books/:id', 'GET', 'Fetch book details'),
    ('Admin', '/books', 'POST', 'Create book (Admin)'),
    ('Orders', '/orders', 'POST', 'Place an order'),
    ('Orders', '/orders/user', 'GET', 'User order history')
]

for feature, endpoint, method, desc in api_data:
    row_cells = table.add_row().cells
    row_cells[0].text = feature
    row_cells[1].text = endpoint
    row_cells[2].text = method
    row_cells[3].text = desc

# 5. Database Schema
doc.add_heading('5. Database Schema', level=1)
doc.add_paragraph("• User: Credentials, Roles (USER/ADMIN).")
doc.add_paragraph("• Book: Title, Price, Description, Inventory.")
doc.add_paragraph("• Category: Classification for books.")
doc.add_paragraph("• Order: Transaction master records.")
doc.add_paragraph("• OrderItem: Selected books in an order.")

# 6. Consumer User Guide
doc.add_heading('6. Consumer User Guide', level=1)
doc.add_heading('6.1 Home & Shop', level=2)
doc.add_picture(os.path.join(artifact_dir, 'consumer_home_1774673615018.png'), width=Inches(6))
doc.add_picture(os.path.join(artifact_dir, 'consumer_shop_1774673628957.png'), width=Inches(6))

doc.add_heading('6.2 Details & Checkout', level=2)
doc.add_picture(os.path.join(artifact_dir, 'consumer_details_1774673644032.png'), width=Inches(6))
doc.add_picture(os.path.join(artifact_dir, 'consumer_cart_1774673662610.png'), width=Inches(6))
doc.add_picture(os.path.join(artifact_dir, 'consumer_checkout_1774673793039.png'), width=Inches(6))
doc.add_picture(os.path.join(artifact_dir, 'consumer_success_1774673844742.png'), width=Inches(6))

# 7. Admin User Guide
doc.add_heading('7. Admin User Guide', level=1)
doc.add_heading('7.1 Dashboard Overview', level=2)
doc.add_picture(os.path.join(artifact_dir, 'admin_dashboard_1774673922220.png'), width=Inches(6))

doc.add_heading('7.2 Management Features', level=2)
doc.add_picture(os.path.join(artifact_dir, 'admin_books_1774673936254.png'), width=Inches(6))
doc.add_picture(os.path.join(artifact_dir, 'admin_categories_1774673949115.png'), width=Inches(6))
doc.add_picture(os.path.join(artifact_dir, 'admin_orders_1774673961576.png'), width=Inches(6))

# Save the document
doc.save(output_file)
print(f"Full Documentation saved to {output_file}")
